"""
Word 文档读取工具

提供 Word 文档的读取和结构提取功能。
使用 python-docx 库。
"""

from pathlib import Path

from docx import Document

from packages.tooling.registry import ToolRegistry


@ToolRegistry.register("read_docx")
def read_docx(path: str) -> dict:
    """
    读取 Word 文档，返回文本内容和结构。

    Args:
        path: Word 文件路径

    Returns:
        {
            "text": "全文文本",
            "structure": {
                "headings": [...],
                "paragraphs": [...],
                "tables": [...]
            },
            "anchors": ["锚点1", "锚点2", ...]
        }
    """
    doc = Document(path)

    # 提取全文文本
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    text = "\n".join(full_text)

    # 提取结构
    structure = extract_docx_structure(path)

    # 提取锚点（模板字段如 {{字段名}}）
    anchors = _extract_anchors(text)

    return {
        "text": text,
        "structure": structure,
        "anchors": anchors,
    }


@ToolRegistry.register("extract_docx_structure")
def extract_docx_structure(path: str) -> dict:
    """
    提取 Word 文档结构。

    Args:
        path: Word 文件路径

    Returns:
        {
            "headings": [
                {"level": 1, "text": "标题1"},
                {"level": 2, "text": "标题2"}
            ],
            "paragraphs": [
                {"style": "Normal", "text": "段落内容"}
            ],
            "tables": [
                {"rows": 3, "cols": 2, "data": [["A1", "A2"], ["B1", "B2"]]}
            ]
        }
    """
    doc = Document(path)

    headings = []
    paragraphs = []

    for para in doc.paragraphs:
        style_name = para.style.name or ""

        # 判断是否是标题
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.replace("Heading ", "").strip())
            except ValueError:
                level = 1
            headings.append({"level": level, "text": para.text})
        elif para.text.strip():
            paragraphs.append({"style": style_name, "text": para.text})

    # 提取表格
    tables = []
    for table in doc.tables:
        rows = len(table.rows)
        cols = len(table.columns)
        data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            data.append(row_data)
        tables.append({"rows": rows, "cols": cols, "data": data})

    return {
        "headings": headings,
        "paragraphs": paragraphs,
        "tables": tables,
    }


def _extract_anchors(text: str) -> list[str]:
    """
    提取模板锚点（{{字段名}} 格式）。

    Args:
        text: 文本内容

    Returns:
        锚点列表，如 ["姓名", "日期", "金额"]
    """
    import re

    pattern = r"\{\{(.+?)\}\}"
    matches = re.findall(pattern, text)
    return list(set(matches))  # 去重
