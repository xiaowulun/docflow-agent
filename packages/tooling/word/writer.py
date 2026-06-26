"""
Word 文档写入工具

提供 Word 文档的写入功能：
- insert_text_by_anchor: 按锚点位置插入文本
- fill_template_fields: 填充模板字段（{{字段名}}）

使用 python-docx 库。
写入时创建新文件，不修改原文件。
"""

import re
import shutil
from pathlib import Path

from docx import Document

from packages.tooling.registry import ToolRegistry


@ToolRegistry.register("insert_text_by_anchor")
def insert_text_by_anchor(path: str, anchor: str, text: str) -> dict:
    """
    按锚点位置插入文本。

    在文档中查找锚点文本，在其后插入新文本。

    Args:
        path: Word 文件路径
        anchor: 锚点文本（要查找的位置）
        text: 要插入的文本

    Returns:
        {
            "success": True,
            "output_path": "新文件路径",
            "message": "操作说明"
        }
    """
    doc = Document(path)
    output_path = _get_output_path(path)

    found = False

    for para in doc.paragraphs:
        if anchor in para.text:
            # 找到锚点，在其后插入文本
            run = para.add_run(f"\n{text}")
            found = True
            break

    if not found:
        return {
            "success": False,
            "output_path": None,
            "message": f"Anchor '{anchor}' not found in document",
        }

    doc.save(output_path)

    return {
        "success": True,
        "output_path": str(output_path),
        "message": f"Text inserted after anchor '{anchor}'",
    }


@ToolRegistry.register("fill_template_fields")
def fill_template_fields(path: str, fields: dict) -> dict:
    """
    填充模板字段。

    将文档中的 {{字段名}} 替换为对应的值。

    Args:
        path: Word 文件路径
        fields: 字段映射，如 {"姓名": "张三", "日期": "2024-01-01"}

    Returns:
        {
            "success": True,
            "output_path": "新文件路径",
            "replaced_count": 5,
            "message": "操作说明"
        }
    """
    doc = Document(path)
    output_path = _get_output_path(path)

    replaced_count = 0

    # 替换段落中的模板字段
    for para in doc.paragraphs:
        new_text, count = _replace_fields(para.text, fields)
        if count > 0:
            _set_paragraph_text(para, new_text)
            replaced_count += count

    # 替换表格中的模板字段
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    new_text, count = _replace_fields(para.text, fields)
                    if count > 0:
                        _set_paragraph_text(para, new_text)
                        replaced_count += count

    doc.save(output_path)

    return {
        "success": True,
        "output_path": str(output_path),
        "replaced_count": replaced_count,
        "message": f"Replaced {replaced_count} template fields",
    }


def _replace_fields(text: str, fields: dict) -> tuple[str, int]:
    """
    替换文本中的模板字段。

    Args:
        text: 原始文本
        fields: 字段映射

    Returns:
        (替换后的文本, 替换次数)
    """
    count = 0

    def replacer(match):
        nonlocal count
        field_name = match.group(1).strip()
        if field_name in fields:
            count += 1
            return str(fields[field_name])
        return match.group(0)  # 未找到字段，保持原样

    new_text = re.sub(r"\{\{(.+?)\}\}", replacer, text)
    return new_text, count


def _set_paragraph_text(paragraph, new_text: str):
    """
    设置段落文本，保留原有格式。

    Args:
        paragraph: 段落对象
        new_text: 新文本
    """
    # 保留第一个 run 的格式
    if paragraph.runs:
        first_run = paragraph.runs[0]
        font = first_run.font

        # 清除所有 run
        for run in paragraph.runs[1:]:
            run.text = ""

        # 设置新文本
        first_run.text = new_text
    else:
        # 没有 run，直接添加
        paragraph.add_run(new_text)


def _get_output_path(original_path: str) -> Path:
    """
    生成输出文件路径。

    在原文件名基础上添加 _filled 后缀。

    Args:
        original_path: 原文件路径

    Returns:
        输出文件路径
    """
    path = Path(original_path)
    stem = path.stem
    suffix = path.suffix
    output_dir = path.parent

    # 避免覆盖，检查文件是否存在
    output_path = output_dir / f"{stem}_filled{suffix}"
    counter = 1
    while output_path.exists():
        output_path = output_dir / f"{stem}_filled_{counter}{suffix}"
        counter += 1

    return output_path
